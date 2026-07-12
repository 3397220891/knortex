import fitz  # PyMuPDF
from docx import Document as DocxDocument
import os
import traceback
from typing import Dict, List


def _join_page_texts(pages: List[str]) -> str:
    """Join per-page text without breaking sentences that span a page boundary.

    A plain "\\n\\n".join() (or the old literal "--- Page N ---" marker) inserts a
    hard break between every page, which severs any sentence PyMuPDF happened to
    split across two pages - fatal for the extractor's phrase-adjacency regexes.
    If the previous page didn't end on sentence-final punctuation, treat the next
    page as a continuation and join with a single space instead.
    """
    joined = ""
    for page_text in pages:
        page_text = page_text.strip()
        if not page_text:
            continue
        if joined and not joined.rstrip().endswith((".", "!", "?", '"', "'")):
            joined = joined.rstrip() + " " + page_text
        else:
            joined = joined + ("\n\n" if joined else "") + page_text
    return joined


class DocumentParser:
    """Parser for extracting text from PDF and Word documents"""

    @staticmethod
    def parse_pdf(file_path: str) -> Dict:
        """Extract text content from PDF files"""
        print(f"Starting PDF parsing: {file_path}")
        
        # Force display all possible error messages
        try:
            print(f"Checking file: {os.path.exists(file_path)}")
            print(f"File size: {os.path.getsize(file_path)}")
            
            # Try opening file directly
            with open(file_path, 'rb') as f:
                header = f.read(10)
                print(f"File header: {header}")
            
            # Use absolute path
            abs_path = os.path.abspath(file_path)
            print(f"Absolute path: {abs_path}")
            
            # Try opening PDF
            print("Opening PDF with fitz...")
            doc = fitz.open(abs_path)
            print(f"PDF opened successfully! Pages: {len(doc)}")
            
            pages = []
            for page_num, page in enumerate(doc):
                page_text = page.get_text()
                pages.append(page_text)
                print(f"Page {page_num + 1}: {len(page_text)} chars")

            page_count = len(doc)
            content = _join_page_texts(pages)

            doc.close()

            return {
                "content": content,
                "page_count": page_count,
                "file_type": "pdf",
                "status": "success"
            }
            
        except Exception as e:
            print(f"PDF parsing FAILED!")
            print(f"Error: {str(e)}")
            print(f"Error type: {type(e).__name__}")
            traceback.print_exc()
            
            return {
                "content": "",
                "page_count": 0,
                "file_type": "pdf", 
                "status": f"error: {str(e)}"
            }
    
    @staticmethod
    def parse_docx(file_path: str) -> Dict:
        """Extract text content from Word documents"""
        print(f"Processing Word document: {os.path.basename(file_path)}")
        
        try:
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # doc.paragraphs skips table cells entirely, so relation-dense content
            # living in tables (org charts, contact lists) was silently dropped.
            table_lines = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        table_lines.append(", ".join(cells))

            return {
                "content": "\n".join(paragraphs + table_lines),
                "paragraph_count": len(paragraphs),
                "file_type": "docx",
                "status": "success"
            }
            
        except Exception as e:
            print(f"DOCX parsing error: {str(e)}")
            return {
                "content": "",
                "paragraph_count": 0,
                "file_type": "docx",
                "status": f"error: {str(e)}"
            }
    @staticmethod
    def parse_txt(file_path: str) -> Dict:
        """Extract text content from plain text files"""
        print(f"Processing text file: {os.path.basename(file_path)}")
    
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        
            return {
                "content": content,
                "file_type": "txt",
                "status": "success"
            }
        
        except Exception as e:
            print(f"TXT parsing error: {str(e)}")
            return {
                "content": "",
                "file_type": "txt", 
                "status": f"error: {str(e)}"
            }