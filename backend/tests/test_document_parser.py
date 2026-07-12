import fitz
from docx import Document as DocxDocument

from services.document_parser import DocumentParser, _join_page_texts


def test_join_page_texts_rejoins_sentence_split_across_a_page_boundary():
    pages = ["TechCorp is headquartered in", "New York."]
    assert _join_page_texts(pages) == "TechCorp is headquartered in New York."


def test_join_page_texts_keeps_distinct_paragraphs_separated():
    pages = ["First page ends cleanly.", "Second page starts a new topic."]
    assert _join_page_texts(pages) == "First page ends cleanly.\n\nSecond page starts a new topic."


def test_join_page_texts_skips_empty_pages():
    pages = ["Some text.", "   ", "More text."]
    assert _join_page_texts(pages) == "Some text.\n\nMore text."


def test_join_page_texts_empty_input():
    assert _join_page_texts([]) == ""


def test_parse_pdf_does_not_break_a_sentence_across_pages(tmp_path):
    doc = fitz.open()
    page1 = doc.new_page()
    page1.insert_text((72, 72), "TechCorp is headquartered in")
    page2 = doc.new_page()
    page2.insert_text((72, 72), "New York.")
    file_path = tmp_path / "sample.pdf"
    doc.save(str(file_path))
    doc.close()

    result = DocumentParser.parse_pdf(str(file_path))

    assert result["status"] == "success"
    assert result["page_count"] == 2
    assert "--- Page" not in result["content"]
    assert "TechCorp is headquartered in New York." in result["content"]


def test_parse_docx_includes_table_content(tmp_path):
    doc = DocxDocument()
    doc.add_paragraph("Company overview follows.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Name"
    table.rows[0].cells[1].text = "Title"
    table.rows[1].cells[0].text = "Jane Doe"
    table.rows[1].cells[1].text = "CTO"
    file_path = tmp_path / "sample.docx"
    doc.save(str(file_path))

    result = DocumentParser.parse_docx(str(file_path))

    assert result["status"] == "success"
    assert "Company overview follows." in result["content"]
    assert "Jane Doe, CTO" in result["content"]
